# from  docx import  Document
# from  docx.shared import  Pt
# from  docx.oxml.ns import  qn
# from docx.shared import Inches
#
# #打开文档
# document = Document()
#
# #加入不同等级的标题
# document.add_heading('Document Title',0)
# document.add_heading(u'二级标题',1)
# document.add_heading(u'二级标题',2)
#
# #添加文本
# paragraph = document.add_paragraph(u'添加了文本')
# #设置字号
# run = paragraph.add_run(u'设置字号')
# run.font.size=Pt(24)
#
# #设置字体
# run = paragraph.add_run('Set Font,')
# run.font.name='Consolas'
#
# #设置中文字体
# run = paragraph.add_run(u'设置中文字体，')
# run.font.name=u'宋体'
# r = run._element
# r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
#
# #设置斜体
# run = paragraph.add_run(u'斜体、')
# run.italic = True
#
# #设置粗体
# run = paragraph.add_run(u'粗体').bold = True
#
# #增加引用
# document.add_paragraph('Intense quote', style='Intense Quote')
#
# #增加有序列表
# document.add_paragraph(
#     u'有序列表元素1',style='List Number'
# )
# document.add_paragraph(
#     u'有序列别元素2',style='List Number'
# )
#
# #增加无序列表
# document.add_paragraph(
#     u'无序列表元素1',style='List Bullet'
# )
# document.add_paragraph(
#     u'无序列表元素2',style='List Bullet'
# )
#
# #增加图片（此处使用相对位置）
# document.add_picture('profile.jpg',width=Inches(1.25))
#
# #增加表格
# table = document.add_table(rows=1,cols=3)
# hdr_cells=table.rows[0].cells
# hdr_cells[0].text="第一列"
# hdr_cells[1].text="第二列"
# hdr_cells[2].text="第三列"
#
# row_cells = table.add_row().cells
# row_cells[0].text = '2'
# row_cells[1].text = '2'
# row_cells[2].text = '2'
#
# # hdr_cells = table.rows[1].cells
# # hdr_cells[0].text = '2'
# # hdr_cells[1].text = 'aerszvfdgx'
# # hdr_cells[2].text = 'abdzfgxfdf'
# #
# # hdr_cells = table.rows[2].cells
# # hdr_cells[0].text = '3'
# # hdr_cells[1].text = 'cafdwvaef'
# # hdr_cells[2].text = 'aabs zfgf'
#
# #增加分页
# document.add_page_break()
#
# #保存文件
# document.save('/Users/skye/Desktop/demo.docx')


# from docx import Document
# from docx.shared import Inches
# import database
#
# data = database.querttestdata()
# print(data)
#
# # python-docx:
# document = Document()
# document.add_heading('Document Title', 0)
# p = document.add_paragraph('A plain paragraph having some ')
# p.add_run('bold').bold = True
# p.add_run(' and some ')
# p.add_run('italic.').italic = True
# document.add_heading('Heading, level 1', level=1)
# document.add_paragraph('Intense quote', style='Intense Quote')
# document.add_paragraph(
#     'first item in unordered list', style='List Bullet'
# )
# document.add_paragraph(
#     'first item in ordered list', style='List Number'
# )
#
# document.add_picture('/Users/skye/Pictures/profile.jpg', width=Inches(1.25))
#
# records = (
#     (3, '101', 'Spam'),
#     (7, '422', 'Eggs'),
#     (4, '631', 'Spam, spam, eggs, and spam')
# )
#
# table = document.add_table(rows=1, cols=5, style='Table Grid')
# hdr_cells = table.rows[0].cells
# hdr_cells[0].text = 'id'
# hdr_cells[1].text = 'data1'
# hdr_cells[2].text = 'data2'
# hdr_cells[3].text = 'data3'
# hdr_cells[4].text = 'data4'
# for tmp in data:
#     row_cells = table.add_row().cells
#     for i in range(len(tmp)):
#         row_cells[i].text = str(tmp[i])
#         # row_cells[1].text = str(tmp[1])
#         # row_cells[2].text = str(tmp[2])
#         # row_cells[3].text = str(tmp[3])
#         # row_cells[4].text = str(tmp[4])
#
# document.add_page_break()
#
# document.save('/Users/skye/Desktop/demo.docx')

from docx import *
from docx.enum.style import WD_STYLE_TYPE
document = Document()
styles = document.styles

# 生成所有表样式
for s in styles:
    if s.type == WD_STYLE_TYPE.TABLE:
        document.add_paragraph("表格样式 :  " + s.name)
        table = document.add_table(3, 3, style=s)
        heading_cells = table.rows[0].cells
        heading_cells[0].text = '第一列内容'
        heading_cells[1].text = '第二列内容'
        heading_cells[2].text = '第三列内容'
        document.add_paragraph("\n")

document.save('/Users/skye/Desktop/demo1.docx')

def splitblank(argument):
    linelen = 30
    ans = ''
    len1 = (linelen - len(argument))/2
    for i in range(0, int(len1)):
        ans += ' '
    ans += argument
    for i in range(0, int(len1)):
        ans += ' '
    if linelen - 2*int(len1) - len(argument) != 0:
        ans += ' '
    ans += '.'
    print(len(ans))
    return ans

# arguments: heading, 使用单位, 生成人, 生成日期, ...
def generateCover():

    document = Document()

    # heading = document.add_heading(level=0)
    heading = document.add_paragraph()
    h1 = heading.add_run('\n\n露天采场矿岩与品位信息报告\n\n\n')
    h1.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
    h1.font.size = Pt(24)  # 一号字体
    h1.bold = True

    # https://www.cnblogs.com/apple2016/p/9635061.html
    # https://blog.csdn.net/wuxiyue238/article/details/81085137
    h1.font.name = u'宋体'
    h1._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Normal'].font.name = u'宋体'  # 修改正文的中文字体类型
    document.styles['Normal'].font.size = Pt(14)  # 修改正文的中文字体类型
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中
    # heading.paragraph_format.left_indent = Inches(0) # 左缩进
    heading.paragraph_format.space_before = Pt(0)  # 上行间距
    # heading.paragraph_format.space_after = Pt(0) # 下行间距
    heading.paragraph_format.line_spacing = 1  # 单倍行距

    p1 = document.add_paragraph()
    p1.add_run('使用单位：').font.size = Pt(14)
    argument1 = 'haha'
    res = splitblank(argument1)
    p1.add_run(res).font.underline = True
    p1.paragraph_format.first_line_indent = Inches(1)
    # p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # p1.paragraph_format.space_after = Pt(0) # 下行间距

    p2 = document.add_paragraph()
    p2.add_run('设备名称：').font.size = Pt(14)
    argument2 = '沈抚支线管道1'
    res = splitblank(argument2)
    p2.add_run(res).font.underline = True
    p2.paragraph_format.first_line_indent = Inches(1)

    p3 = document.add_paragraph()
    p3.add_run('设备类别：').font.size = Pt(14)
    argument3 = '长输管道'
    res = splitblank(argument3)
    p3.add_run(res).font.underline = True
    p3.paragraph_format.first_line_indent = Inches(1)

    p4 = document.add_paragraph()
    p4.add_run('设备品种：').font.size = Pt(14)
    argument4 = '天然气管道'
    res = splitblank(argument4)
    p4.add_run(res).font.underline = True
    p4.paragraph_format.first_line_indent = Inches(1)

    p5 = document.add_paragraph()
    p5.add_run('压力管道代码：').font.size = Pt(14)
    argument5 = '81201'
    res = splitblank(argument5)
    p5.add_run(res).font.underline = True
    p5.paragraph_format.first_line_indent = Inches(1)

    p6 = document.add_paragraph()
    p6.add_run('检验日期：').font.size = Pt(14)
    argument6 = '2019.04.03'
    res = splitblank(argument6)
    p6.add_run(res).font.underline = True
    p6.paragraph_format.first_line_indent = Inches(1)

    p7 = document.add_paragraph('\n\n\n\n\n')

    p8 = document.add_paragraph()
    p8.add_run('东北大学软件学院').font.size = Pt(14)
    p8.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中

    document.save('/Users/skye/Desktop/demo1.docx')