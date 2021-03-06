from docx import Document
from docx.shared import Inches, RGBColor, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn

def generateCover(document, doc_heading = "", doc_company = "", doc_code = "", doc_type = "", doc_reporter = "", doc_date = ""):
    print("generate Cover")
    # document = Document()
    # heading = document.add_heading(level=0)
    heading = document.add_paragraph()
    doc_heading = '\n\n' + doc_heading + '\n\n\n'
    h1 = heading.add_run(doc_heading)
    h1.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
    h1.font.size = Pt(26)  # 一号字体
    h1.bold = True

    # https://www.cnblogs.com/apple2016/p/9635061.html
    # https://blog.csdn.net/wuxiyue238/article/details/81085137
    h1.font.name = u'宋体'
    h1._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Normal'].font.name = u'宋体'  # 修改正文的中文字体类型
    document.styles['Normal'].font.size = Pt(12)  # 修改正文的中文字体类型
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中
    # heading.paragraph_format.left_indent = Inches(0) # 左缩进
    heading.paragraph_format.space_before = Pt(0)  # 上行间距
    # heading.paragraph_format.space_after = Pt(0) # 下行间距
    heading.paragraph_format.line_spacing = 1  # 单倍行距

    row = 5
    column = 2
    table = document.add_table(rows=row, cols=column, style='Normal Table')

    row0 = table.rows[0]
    p1 = row0.cells[0].paragraphs[0]
    p1.add_run('使用单位：').font.size = Pt(15)
    p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2 = row0.cells[1].paragraphs[0]
    p2.add_run(doc_company).font.size = Pt(15)
    p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    row0 = table.rows[1]
    p1 = row0.cells[0].paragraphs[0]
    p1.add_run('报告编号：').font.size = Pt(15)
    p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2 = row0.cells[1].paragraphs[0]
    p2.add_run(doc_code).font.size = Pt(15)
    p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    row0 = table.rows[2]
    p1 = row0.cells[0].paragraphs[0]
    p1.add_run('报告类型：').font.size = Pt(15)
    p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2 = row0.cells[1].paragraphs[0]
    p2.add_run(doc_type).font.size = Pt(15)
    p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    row0 = table.rows[3]
    p1 = row0.cells[0].paragraphs[0]
    p1.add_run('报告人员：').font.size = Pt(15)
    p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2 = row0.cells[1].paragraphs[0]
    p2.add_run(doc_reporter).font.size = Pt(15)
    p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    row0 = table.rows[4]
    p1 = row0.cells[0].paragraphs[0]
    p1.add_run('报告日期：').font.size = Pt(15)
    p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2 = row0.cells[1].paragraphs[0]
    p2.add_run(doc_date).font.size = Pt(15)
    p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for i in range(0, row):
        # row0 = table.rows[i]
        # p1 = row0.cells[0].paragraphs[0]
        # p1.add_run('使用单位：')
        # p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # table.cell(i, 0).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        # p2 = row0.cells[1].paragraphs[0]
        # p2.add_run('haha')
        # p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # table.cell(i, 1).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        table.rows[i].height = Cm(1.5)
        # table.cell(i, 0).width = Cm(6)

    p7 = document.add_paragraph('\n\n\n\n')
    p8 = document.add_paragraph()
    p8.add_run(doc_company).font.size = Pt(15)
    p8.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中

    document.add_page_break()
    # document.save('/Users/skye/Desktop/demo1.docx')

def generateCatalog(document, catalog, catalogContent, doc_code = "", doc_heading = ""):
    print("generate Catalog")
    # document = Document()
    heading = document.add_paragraph()
    doc_heading = doc_heading + "目录"
    h1 = heading.add_run(doc_heading)
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.font.size = Pt(16)
    h1.bold = True
    h1.font.name = u'宋体'
    h1._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Normal'].font.name = u'宋体'  # 修改正文的中文字体类型
    document.styles['Normal'].font.size = Pt(12)  # 修改正文的中文字体类型
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中
    # heading.paragraph_format.left_indent = Inches(0) # 左缩进
    heading.paragraph_format.space_before = Pt(0)  # 上行间距
    # heading.paragraph_format.space_after = Pt(0) # 下行间距
    heading.paragraph_format.line_spacing = 1  # 单倍行距

    p = document.add_paragraph()
    p.add_run("报告编号: "+doc_code).font.size = Pt(14)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    row = 11
    column = 4
    table = document.add_table(rows=row, cols=column, style='Table Grid')
    row0 = table.rows[0]
    p1 = row0.cells[0].paragraphs[0]
    p1.add_run('序号')
    p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2 = row0.cells[1].paragraphs[0]
    p2.add_run('报告内容')
    p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2 = row0.cells[2].paragraphs[0]
    p2.add_run('页码')
    p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2 = row0.cells[3].paragraphs[0]
    p2.add_run('附页')
    p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i in range(1, row):
        row0 = table.rows[i]
        p1 = row0.cells[0].paragraphs[0]
        p1.add_run((str)(i))
        p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2 = row0.cells[1].paragraphs[0]
        if i in catalog:
            p2.add_run(catalogContent[i-1])
        else:
            p2.add_run("/")
        p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p3 = row0.cells[2].paragraphs[0]
        p3.add_run('/')
        p3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p4 = row0.cells[3].paragraphs[0]
        p4.add_run('/')
        p4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i in range(0, row):
        table.rows[i].height = Cm(1.3)
        for j in range(0, column):
            table.cell(i, j).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if j == 0:
                table.cell(i, 0).width = Cm(4)
            elif j == 1:
                table.cell(i, 1).width = Cm(10)
            elif j == 2: table.cell(i, 2).width = Cm(4)
            else: table.cell(i, 3).width = Cm(4)

    document.add_page_break()
    # document.save('/Users/skye/Desktop/directory.docx')

def generateContent1(document, arguments, doc_code="", doc_content1=""):
    print("generate Content1")
    heading = document.add_paragraph()
    h1 = heading.add_run(doc_content1)
    # h1 = heading.add_run("1.资料审查报告")
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.font.size = Pt(16)
    h1.bold = True
    h1.font.name = u'宋体'
    h1._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal'].font.size = Pt(12)
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.line_spacing = 1
    p = document.add_paragraph()
    p.add_run("报告编号: " + doc_code).font.size = Pt(14)
    # p.add_run("报告编号: " + "20180403-A").font.size = Pt(12)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    table1 = document.add_table(rows=7, cols=4, style='Table Grid')
    table2 = document.add_table(rows=2, cols=2, style='Table Grid')
    table3 = document.add_table(rows=1, cols=2, style='Table Grid')

    tmp = ["矿区名称", "矿区编号", "管理单位", "单位编号", "操作人", "操作类型", "原始数据(个)", "新增数据(个)", "误差(更新前)", "准确率(更新前)",
           "误差(更新后)", "准确率(更新后)", "上次检验日期", "上次检验记录编号", "原始资料审查问题记载", "历次定期检验问题记载", "审核日期(年/月/日)"]
    for i in range(0, 7):
        row = table1.rows[i]
        for j in range(0, 2):
            p = row.cells[j*2].paragraphs[0]
            p.add_run(tmp[i*2+j])
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    row = table2.rows[0]
    p = row.cells[0].paragraphs[0]
    p.add_run(tmp[14])
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    row = table2.rows[1]
    p = row.cells[0].paragraphs[0]
    p.add_run(tmp[15])
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    row = table3.rows[0]
    p = row.cells[0].paragraphs[0]
    p.add_run(tmp[16])
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 根据arguments生成文档
    count = 0
    for i in range(7):
        row = table1.rows[i]
        for j in [1, 3]:
            p = row.cells[j].paragraphs[0]
            p.add_run(arguments[count])
            count += 1
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    row = table2.rows[0]
    p = row.cells[1].paragraphs[0]
    p.add_run(arguments[count])
    count += 1
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    row = table2.rows[1]
    p = row.cells[1].paragraphs[0]
    p.add_run(arguments[count])
    count += 1
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    row = table3.rows[0]
    p = row.cells[1].paragraphs[0]
    p.add_run(arguments[count])
    count += 1
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i in range(0, 7):
        table1.rows[i].height = Cm(1.8)
        for j in range(0, 4):
            table1.cell(i, j).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if j == 0 or j == 2:
                table1.cell(i, j).width = Cm(5)
            else:
                table1.cell(i, j).width = Cm(7)
    for i in range(0, 2):
        table2.rows[i].height = Cm(1.8)
        table2.cell(i, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table2.cell(i, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table3.rows[0].height = Cm(1.8)
    table3.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table3.cell(0, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    document.add_page_break()

def generateContent2(document, arguments, doc_content2=""):
    print("generate Content2")
    print(arguments)
    heading = document.add_paragraph()
    h1 = heading.add_run(doc_content2)
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.font.size = Pt(16)
    h1.bold = True
    h1.font.name = u'宋体'
    h1._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal'].font.size = Pt(12)
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.line_spacing = 1
    # doc_2_text
    p = document.add_paragraph(arguments[2])
    p.paragraph_format.space_after = 0
    p = document.add_paragraph("\n")
    p.paragraph_format.space_after = 0

    if len(arguments[1]) == 0:
        pass
    else:
        # doc_2_table
        row = len(arguments[1]) + 1
        column = 5
        table = document.add_table(rows=row, cols=column, style='Table Grid')
        row0 = table.rows[0]
        p1 = row0.cells[0].paragraphs[0]
        p1.add_run('样本编号')
        p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2 = row0.cells[1].paragraphs[0]
        p2.add_run('x')
        p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2 = row0.cells[2].paragraphs[0]
        p2.add_run('y')
        p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2 = row0.cells[3].paragraphs[0]
        p2.add_run('z')
        p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2 = row0.cells[4].paragraphs[0]
        p2.add_run('所属类别')
        p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            import pandas as pd
            import helper
            df = pd.read_csv(helper.data_path, header=None)
            dataset = df.values[0:5][:, 1:].transpose()[:, [0, 2, 3, 4, 1]]
            rowNum = dataset.shape[0]
        except:
            print("generateContent2 catch exception: 读取文件出错")
        rownum = 1
        for i in arguments[1]:
            if i < rowNum:
                row0 = table.rows[rownum]
                rownum += 1
                for j in range(5):
                    p1 = row0.cells[j].paragraphs[0]
                    p1.add_run(str(dataset[i][j]))
                    p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for i in range(0, row):
            if i == 0:
                table.rows[i].height = Cm(1.5)
            else:
                table.rows[i].height = Cm(0.5)
            for j in range(0, column):
                table.cell(i, j).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                if j == 0:
                    table.cell(i, 0).width = Cm(4)
                else:
                    # print(table.cell(i, j).width)
                    table.cell(i, j).width = Cm(3)

    # doc_2_tablename
    p = document.add_paragraph("表2.1 "+arguments[0])
    p.paragraph_format.space_after = 0
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph("\n")
    p.paragraph_format.space_after = 0

    document.add_page_break()

def generateContent3(document, arguments, doc_content3=""):
    print("generate Content3")
    heading = document.add_paragraph()
    h1 = heading.add_run(doc_content3)
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.font.size = Pt(16)
    h1.bold = True
    h1.font.name = u'宋体'
    h1._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal'].font.size = Pt(12)
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.line_spacing = 1
    # doc_3_text
    p = document.add_paragraph(arguments[2])
    p.paragraph_format.space_after = 0
    p = document.add_paragraph("\n")
    p.paragraph_format.space_after = 0

    # doc_3_imgpath, doc_3_imgname
    imgnum = len(arguments[0])
    if imgnum == 0:
        pass
    else:
        for i in range(imgnum):
            table = document.add_table(rows=1, cols=1, style='Normal Table')
            cell = table.rows[0].cells[0]
            p = cell.paragraphs[0]
            run = p.add_run()
            run.add_picture(arguments[0][i], width=Inches(2.5))
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            p = document.add_paragraph("图3.1 " + arguments[1][0])
            p.paragraph_format.space_after = 0
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # p = document.add_paragraph("\n")
            # p.paragraph_format.space_after = 0
    document.add_page_break()

def generateContent4(document, doc_4_text, doc_content4=""):
    print("generate Content4")
    heading = document.add_paragraph()
    h1 = heading.add_run(doc_content4)
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.font.size = Pt(16)
    h1.bold = True
    h1.font.name = u'宋体'
    h1._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal'].font.size = Pt(12)
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.line_spacing = 1
    # doc_4_text
    p = document.add_paragraph(doc_4_text)
    p.paragraph_format.space_after = 0

def generatedoc(document, path):
    document.save(path)
